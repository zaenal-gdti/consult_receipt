from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
from docx.shared import Pt
import pandas as pd
import os
import shutil
from tqdm import tqdm

def row_to_pdf(data):
    f = open('template.docx', 'rb')
    doc = Document(f)
    
    name = data['Name']
    consult_on = pd.to_datetime(data['Consultation Date']).strftime('%d-%b-%Y')
    user_id = data['Card Number']#'0012'
    doctor = data['Doctor Name'] 
    cons_time = data['Start Time']#.astype(str)
    comp = data['Corporate Name']
    cons_id = data['Consultation ID']
    claim_id = data['Claim ID']
    icd = data['ICDX']+ ' ' + data['Diagnosis']
    
    
    replace_word = {'Patient Name : ': f'Patient Name: {name}', 
                    'Patient User ID : ': f'Patient User ID : {user_id}',
                    'Doctor Name\t\t: ' : f'Doctor Name\t\t: {doctor}',
                    'Consult Time\t\t: ' : f'Consult Time\t\t: {cons_time}',
                    'Company Name\t: ' : f'Company Name\t: {comp}',
                    'Consult ID\t\t: ' : f'Consult ID\t\t: {cons_id}',
                    'Claim ID\t\t: ': f'Claim ID\t\t: {claim_id}',
                    'ICDX\t\t\t: ' : f'ICDX\t\t\t: {icd}'
    }

    
    tbl1 = {'consult_price' :  data['Consult Fee'],
            'drug_price' : data['Rx Fee']
            }
    tbl1['total_fare'] = tbl1['consult_price'] + tbl1['drug_price']

    
    
    

    for word in replace_word:
        for p in doc.paragraphs:
            if p.text.find(word) >= 0:
                p.text = p.text.replace(word, replace_word[word])
    
    par2 = doc.paragraphs[2]
    
    par2.text = f'Consulted on : {consult_on}'
    
    par2.runs[0].bold = True


    tbl_val1 = tbl1['consult_price']
    tbl_val2 = tbl1['drug_price']
    tbl_val3 = tbl1['total_fare']
    
    
    
    doc.tables[0].cell(0, 1).text = f'Rp {tbl_val1}'
    doc.tables[0].cell(0, 1).paragraphs[0].runs[0].bold = True
    doc.tables[0].cell(0, 1).paragraphs[0].alignment = 2
    
    doc.tables[0].cell(1, 1).text = f'Rp {tbl_val2}'
    doc.tables[0].cell(1, 1).paragraphs[0].runs[0].bold = True
    doc.tables[0].cell(1, 1).paragraphs[0].alignment = 2
    
    doc.tables[0].cell(2, 1).text = f'Rp {tbl_val3}'
    doc.tables[0].cell(2, 1).paragraphs[0].runs[0].bold = True
    doc.tables[0].cell(2, 1).paragraphs[0].alignment = 2
    doc.tables[0].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(18)
    
    for i in range(3):
        doc.tables[0].cell(i, 1).paragraphs[0].paragraph_format.space_after = Pt(0)

    obat = data['pres_all']
    if obat != 'nan':
    
        
        
        obat_idn = obat.split('|')
        for i in range(len(obat_idn)):
            obat_cl = obat_idn[i].split(';')
            j = i + 1
            doc.tables[1].cell(j, 0).text = obat_cl[0]
            doc.tables[1].cell(j, 1).text = obat_cl[2].replace('.0', '')
            doc.tables[1].cell(j, 2).text = obat_cl[1].replace('.0', '')
            doc.tables[1].cell(j, 3).text = obat_cl[3].replace('.0', '')
    
            doc.tables[1].cell(j, 0).paragraphs[0].runs[0].font.size = Pt(8)
            doc.tables[1].cell(j, 1).paragraphs[0].runs[0].font.size = Pt(8)
            doc.tables[1].cell(j, 2).paragraphs[0].runs[0].font.size = Pt(8)
            doc.tables[1].cell(j, 3).paragraphs[0].runs[0].font.size = Pt(8)
        
            for k in range(4):
                doc.tables[1].cell(j, k).paragraphs[0].paragraph_format.space_after = Pt(0)
        
    doc.save(f'output/{consult_on}_Consultation_Receipt_{name}.docx')
    f.close()

    # subprocess.run(['libreoffice', '--convert-to', 'pdf' ,
    #                 f'output/{consult_on}_Consultation_Receipt_{name}.docx', '--outdir', 'output/']
    #                ,stdout=subprocess.DEVNULL,
    #                 stderr=subprocess.STDOUT
    #             )

    # old_doc = [x for x in os.listdir('output') if 'docx' in x]
    # for i in old_doc:
    #     os.remove(f'output/{i}')


def mail_merge(file):
    shutil.rmtree('output')
    os.mkdir('output')
    
    df_input = pd.read_excel('Input/sample_consult_file_w_prescription.xlsx')
    df_input['ICDX'] = df_input.filter(like = 'ICD').astype(str).agg(','.join, axis=1).str.replace(',nan', '')
    df_input['Diagnosis'] = df_input.filter(like = 'Diagnosis ').astype(str).agg(','.join, axis=1).str.replace(',nan', '')
    max_pres = df_input.filter(like = 'obat_').columns[-1].split('_')[1]

    df_input['pres_all'] = ''
    for i in range(int(max_pres)):
        j = i + 1
        new_pres = df_input[[x +'_'+ str(j) for x in ['obat', 'harga', 'jumlah', 'total']]].astype(str).agg(';'.join, axis=1)
        df_input['pres_all'] = df_input['pres_all'] + '|' + new_pres
    # for i in 
    # df_input.filter(like = 'obat_')
    df_input['pres_all'] = df_input['pres_all'].str.replace('^.', '', regex = True)
    df_input['pres_all']  = df_input['pres_all'].str.replace('nan;', '').str.replace('nan|', '').str.replace('|nan', '')

    df_input['Rx Fee'] = df_input['Rx Fee'].fillna(0)
    df_input['Consult Fee'] = df_input['Consult Fee'].fillna(0)
    for index, row in tqdm(df_input.iterrows()):
        row_to_pdf(row)